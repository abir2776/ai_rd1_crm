"""
step 1: open CV file
step 2: extract text from CV
step 3: remove adverb, conjunction, Interjections and articles from text
Step 4: remove special characters
Step 5: Remove single character words
Step 6: return cleaned text
"""

import os
import re
from typing import List


def process_cv(file_path: str) -> str:
    """
    step 1: open CV file (accepts only .pdf or .docx) and read
    step 2: extract text from CV
    step 3: remove adverbs, conjunctions, interjections, and articles from text
    step 4: remove special characters
    step 5: split text into words based on space
    step 6: return list of words
    """
    # ---- step 1 & 2: read file and extract raw text ----
    ext = file_path.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        try:
            from PyPDF2 import PdfReader
        except ImportError as e:
            raise ImportError(
                "PyPDF2 is required to read PDFs. Install with: pip install PyPDF2"
            ) from e

        reader = PdfReader(file_path)
        texts = []
        for page in reader.pages:
            # Some PDFs return None for extract_text(); guard against that
            page_text = page.extract_text() or ""
            texts.append(page_text)
        text = "\n".join(texts)

    elif ext == "docx":
        try:
            import docx  # python-docx
        except ImportError as e:
            raise ImportError(
                "python-docx is required to read DOCX. Install with: pip install python-docx"
            ) from e

        doc = docx.Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs)
    else:
        raise ValueError("Unsupported file format. Only .pdf and .docx are accepted.")

    # Normalize to lowercase for consistent matching
    text = text.lower()

    # ---- step 3: remove adverbs, conjunctions, interjections, and articles ----
    # NOTE: Without a POS tagger, we approximate using curated lists of common words in these categories.
    # You can expand these sets as needed.
    articles = {"a", "an", "the"}
    conjunctions = {
        "and",
        "but",
        "or",
        "nor",
        "for",
        "so",
        "yet",
        "although",
        "though",
        "because",
        "since",
        "unless",
        "while",
        "whereas",
        "either",
        "neither",
        "both",
        "whether",
        "after",
        "before",
        "once",
        "till",
        "until",
        "when",
        "whenever",
        "as",
        "if",
        "than",
    }
    interjections = {
        "oh",
        "ah",
        "wow",
        "hey",
        "alas",
        "oops",
        "uh",
        "um",
        "hmm",
        "yikes",
        "bravo",
        "eureka",
    }
    # Common adverbs (non-exhaustive)
    adverbs = {
        "very",
        "really",
        "just",
        "also",
        "however",
        "therefore",
        "moreover",
        "then",
        "too",
        "quite",
        "rather",
        "fairly",
        "nearly",
        "almost",
        "already",
        "still",
        "yet",
        "ever",
        "never",
        "always",
        "sometimes",
        "often",
        "usually",
        "rarely",
        "seldom",
        "hardly",
        "barely",
        "truly",
        "clearly",
        "simply",
        "highly",
        "extremely",
        "particularly",
        "typically",
        "generally",
        "approximately",
        "about",
        "around",
        "roughly",
        "specifically",
        "finally",
        "additionally",
        "besides",
        "thus",
        "hence",
        "furthermore",
        "likewise",
    }

    prepositions = {
        "in",
        "on",
        "at",
        "by",
        "for",
        "with",
        "about",
        "against",
        "between",
        "into",
        "through",
        "during",
        "before",
        "after",
        "above",
        "below",
        "to",
        "from",
        "up",
        "down",
        "of",
        "off",
        "over",
        "under",
        "within",
        "without",
        "along",
        "across",
        "behind",
        "beyond",
        "near",
        "since",
        "until",
    }

    adjectives = {
        "quick",
        "lazy",
        "happy",
        "sad",
        "bright",
        "dark",
        "excellent",
        "poor",
        "strong",
        "weak",
    }

    pronouns = {
        "i",
        "you",
        "he",
        "she",
        "it",
        "we",
        "they",
        "me",
        "him",
        "her",
        "us",
        "them",
    }

    auxiliary_verbs = {
        "am",
        "is",
        "are",
        "was",
        "were",
        "be",
        "being",
        "been",
        "have",
        "has",
        "had",
        "do",
        "does",
        "did",
        "will",
        "shall",
        "would",
        "should",
        "may",
        "might",
        "must",
    }

    unwanted_words = {
        # Add any additional unwanted words here
        "cv",
        "resume",
        "curriculum",
        "vitae",
        "hopefully",
        "kindly",
        "regards",
        "sincerely",
        "thank you",
        "please",
        "contact",
        "information",
        "details",
        "profile",
        "summary",
        "objective",
        "experience",
        "skills",
        "education",
        "certifications",
        "achievements",
        "references",
        "hobbies",
        "interests",
        "linkedin",
        "github",
        "portfolio",
        "address",
        "phone",
        "email",
        "website",
        "mobile",
        "gmail",
        "outlook",
        "yahoo",
        "hotmail",
        "com",
        "now",
        "today",
        "currently",
        "previously",
        "formerly",
        "recently",
        "soon",
        "later",
        "quickly",
        "slowly",
        "late",
        "easily",
        "hard",
        "difficult",
        "simple",
        "complex",
        "various",
        "multiple",
        "several",
        "numerous",
        "different",
        "diverse",
        "varied",
        "many",
        "few",
        "some",
        "any",
        "all",
        "most",
        "least",
        "variety",
        "assorted",
        "miscellaneous",
        "links",
        "follow",
        "followers",
        "following",
        "connect",
        "network",
        "social",
        "media",
        "overview",
        "background",
        "history",
        "timeline",
        "milestones",
        "progress",
        "development",
        "growth",
        "improvement",
        "advancement",
        "evolution",
        "expansion",
        "extension",
        "enhancement",
        "upgrade",
        "update",
        "revision",
        "modification",
        "adjustment",
        "change",
        "exiting",
        "leaving",
        "departing",
        "relocating",
        "transitioning",
        "shifting",
        "switching",
        "altering",
        "entering",
        "transforming",
        "converting",
        "adapting",
        "adjusting",
        "modifying",
        "customizing",
        "tailoring",
        "personalizing",
        "specializing",
        "focusing",
        "concentrating",
        "dedicating",
        "committing",
        "investing",
        "allocating",
        "devoting",
        "linkedin",
        "gitlab",
        "bitbucket",
        "stackoverflow",
        "facebook",
        "twitter",
        "instagram",
        "tiktok",
        "snapchat",
        "pinterest",
        "reddit",
        "quora",
        "medium",
        "tumblr",
        "flickr",
        "vimeo",
        "youtube",
        "behance",
        "dribbble",
        "slideshare",
        "inlinkedin",
        "githubgithub",
        "linkedinlinkedin",
    }

    remove_words = (
        articles
        | conjunctions
        | interjections
        | adverbs
        | prepositions
        | adjectives
        | pronouns
        | auxiliary_verbs
        | unwanted_words
    )
    # Build a regex that removes whole-word matches followed by a single space to avoid double spaces later;
    # we'll clean spaces again after special-character removal.
    if remove_words:
        pattern = r"\b(?:%s)\b" % "|".join(
            map(re.escape, sorted(remove_words, key=len, reverse=True))
        )
        text = re.sub(pattern, " ", text, flags=re.IGNORECASE)

    # ---- step 4: remove special characters (keep letters, digits, and spaces) ----
    text = re.sub(r"[^a-z0-9\s]", " ", text)

    # Collapse multiple whitespace to single spaces
    text = re.sub(r"\s+", " ", text).strip()

    # Collapse newlines to spaces
    # text = re.sub(r"\n+", " ", text).strip()

    # ---- step 5: Remove single character words ----
    text = re.sub(r"\b\w\b", " ", text)
    text = re.sub(r"\s+", " ", text).strip()

    # ---- step 6: split text into words based on space ----
    # words = text.split(" ") if text else []

    # # Remove any residual empty strings (paranoia)
    # words = [w for w in words if w]

    # ---- step 6: return cleaned text ----
    return {"total_character": len(text), "cleaned_text": text.strip()}


def process_cv(file_path: str) -> str:
    """
    step 1: open CV file (accepts only .pdf or .docx) and read
    step 2: extract text from CV
    step 3: remove adverbs, conjunctions, interjections, and articles from text
    step 4: remove special characters
    step 5: split text into words based on space
    step 6: return list of words
    """
    # ---- step 1 & 2: read file and extract raw text ----
    ext = file_path.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        try:
            from PyPDF2 import PdfReader
        except ImportError as e:
            raise ImportError(
                "PyPDF2 is required to read PDFs. Install with: pip install PyPDF2"
            ) from e

        reader = PdfReader(file_path)
        texts = []
        for page in reader.pages:
            # Some PDFs return None for extract_text(); guard against that
            page_text = page.extract_text() or ""
            texts.append(page_text)
        text = "\n".join(texts)

    elif ext == "docx":
        try:
            import docx  # python-docx
        except ImportError as e:
            raise ImportError(
                "python-docx is required to read DOCX. Install with: pip install python-docx"
            ) from e

        doc = docx.Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs)
    else:
        raise ValueError("Unsupported file format. Only .pdf and .docx are accepted.")

    # Normalize to lowercase for consistent matching
    text = text.lower()

    # ---- step 4: remove special characters (keep letters, digits, and spaces) ----
    text = re.sub(r"[^a-z0-9\s]", " ", text)

    # Collapse multiple whitespace to single spaces
    text = re.sub(r"\s+", " ", text).strip()

    # Collapse newlines to spaces
    # text = re.sub(r"\n+", " ", text).strip()

    # ---- step 5: Remove single character words ----
    # text = re.sub(r"\b\w\b", " ", text)
    # text = re.sub(r"\s+", " ", text).strip()

    # ---- step 6: split text into words based on space ----
    # words = text.split(" ") if text else []

    # # Remove any residual empty strings (paranoia)
    # words = [w for w in words if w]

    # ---- step 6: return cleaned text ----
    return {"total_character": len(text), "cleaned_text": text.strip()}


##### For testing purposes only #####

# # Example (remove or adapt in your codebase):
# result = process_cv("ai_candidate_skill/Kawsar_s_CV.pdf")
# print(result)
